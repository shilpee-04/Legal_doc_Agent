# mcp_tools/doc_parser.py
from dotenv import load_dotenv
load_dotenv()

import json
import pypdf
import docx
from mcp.server.fastmcp import FastMCP

mcp = FastMCP("doc-parser")

def _parse_pdf(file_path: str) -> str:
    reader = pypdf.PdfReader(file_path)
    text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text.strip()

def _parse_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    return text.strip()

@mcp.tool()
def parse_document(file_path: str, file_type: str) -> str:
    """Extract full text from a PDF or DOCX file.
    
    Args:
        file_path: Absolute path to the file
        file_type: Type of file - either 'pdf' or 'docx'
    """
    try:
        if file_type == "pdf":
            return _parse_pdf(file_path)
        elif file_type == "docx":
            return _parse_docx(file_path)
        else:
            return f"Error: Unsupported file type '{file_type}'. Use 'pdf' or 'docx'."
    except FileNotFoundError:
        return f"Error: File not found at {file_path}"
    except Exception as e:
        return f"Error: {str(e)}"

@mcp.tool()
def get_document_info(file_path: str, file_type: str) -> str:
    """Get metadata about a document like page count and word count.
    
    Args:
        file_path: Absolute path to the file
        file_type: Type of file - either 'pdf' or 'docx'
    """
    try:
        if file_type == "pdf":
            reader = pypdf.PdfReader(file_path)
            text = _parse_pdf(file_path)
            info = {
                "page_count": len(reader.pages),
                "word_count": len(text.split()),
                "char_count": len(text)
            }
        elif file_type == "docx":
            doc = docx.Document(file_path)
            text = _parse_docx(file_path)
            info = {
                "paragraph_count": len(doc.paragraphs),
                "word_count": len(text.split()),
                "char_count": len(text)
            }
        else:
            return f"Error: Unsupported file type '{file_type}'"
        
        return json.dumps(info, indent=2)
    except FileNotFoundError:
        return f"Error: File not found at {file_path}"
    except Exception as e:
        return f"Error: {str(e)}"

if __name__ == "__main__":
    print("MCP doc-parser server starting...")
    mcp.run(transport="stdio")